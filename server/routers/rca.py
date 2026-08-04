"""RCA router — generate root cause analysis + DOCX export."""
from io import BytesIO
from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from ..security.jwt import get_current_user
from ..services.rag_service import rag

router = APIRouter(prefix="/api/rca", tags=["rca"])


@router.post("/{tag}")
def generate_rca(tag: str, user: dict = Depends(get_current_user)):
    evidence = rag.asset_evidence(tag)
    return rag.generate_rca(tag, evidence)


@router.get("/export/{tag}")
def export_rca_docx(tag: str, user: dict = Depends(get_current_user)):
    evidence = rag.asset_evidence(tag)
    rca = rag.generate_rca(tag, evidence)

    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        return {"error": "python-docx not installed"}

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    doc.add_heading(f"Root Cause Analysis — {rca['assetTag']}", level=0)
    doc.add_paragraph(f"Confidence: {rca['confidence']:.0%}")

    doc.add_heading("Observed Problem", level=1)
    doc.add_paragraph(rca["observedProblem"])

    doc.add_heading("Probable Causes", level=1)
    for c in rca["probableCauses"]:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("Recorded Corrective Actions", level=1)
    for a in rca["recordedActions"]:
        doc.add_paragraph(a, style="List Bullet")

    if rca["measurements"]:
        doc.add_heading("Relevant Measurements", level=1)
        for m in rca["measurements"]:
            doc.add_paragraph(m, style="List Bullet")

    if rca["eventDates"]:
        doc.add_heading("Event Dates", level=1)
        doc.add_paragraph(", ".join(rca["eventDates"]))

    doc.add_heading("Recommended Investigation", level=1)
    for r in rca["recommendedInvestigation"]:
        doc.add_paragraph(r, style="List Number")

    doc.add_heading("Preventive Actions", level=1)
    for p in rca["preventiveActions"]:
        doc.add_paragraph(p, style="List Number")

    doc.add_paragraph("")
    doc.add_paragraph(rca["disclaimer"]).italic = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=RCA-{rca['assetTag']}.docx"},
    )
