import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_rca_docx(rca_data: dict) -> bytes:
    doc = Document()
    
    # Title
    title = doc.add_heading("IndusMind AI — Root Cause Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle / Asset info
    tag = rca_data.get("assetTag", "UNKNOWN")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run(f"Asset Tag: {tag} | Confidence: {int(rca_data.get('confidence', 0) * 100)}%")
    run_sub.bold = True
    run_sub.font.size = Pt(13)
    
    doc.add_paragraph() # spacing

    # Observed Problem
    doc.add_heading("1. Observed Problem", level=1)
    doc.add_paragraph(rca_data.get("observedProblem", "N/A"))

    # Probable Causes
    doc.add_heading("2. Probable Causes", level=1)
    for cause in rca_data.get("probableCauses", []):
        doc.add_paragraph(cause, style="List Bullet")

    # Recommended Investigation
    doc.add_heading("3. Recommended Investigation Steps", level=1)
    for step in rca_data.get("recommendedInvestigation", []):
        doc.add_paragraph(step, style="List Number")

    # Preventive Actions
    doc.add_heading("4. Preventive Actions", level=1)
    for action in rca_data.get("preventiveActions", []):
        doc.add_paragraph(action, style="List Bullet")

    # Timeline & Event Dates
    dates = rca_data.get("eventDates", [])
    if dates:
        doc.add_heading("5. Related Event Dates", level=1)
        doc.add_paragraph(", ".join(dates))

    # Citations
    citations = rca_data.get("citations", [])
    if citations:
        doc.add_heading("6. Evidence Citations", level=1)
        for c in citations:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{c.get('source', '')} (p. {c.get('page', '')})]: ").bold = True
            p.add_run(c.get("excerpt", ""))

    # Disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    run_disc = disclaimer.add_run(rca_data.get("disclaimer", ""))
    run_disc.italic = True
    run_disc.font.size = Pt(9)
    run_disc.font.color.rgb = RGBColor(120, 120, 120)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
